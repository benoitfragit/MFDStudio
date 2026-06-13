#!/usr/bin/env python3
from __future__ import annotations

import argparse
import datetime as dt
import re
import urllib.request
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor


BODY_TEXT_COLOR = RGBColor(0x1F, 0x29, 0x33)
TABLE_HEADER_TEXT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)
CHAPTER_GROUP_COLORS = {
    "blue": "1E88E5",
    "purple": "6A1B9A",
    "orange": "E65100",
    "teal": "00838F",
}
CALLOUT_STYLES = {
    "NOTE": {"icon": "ℹ", "fill": "EEF6FF", "label_fill": "DCEBFF", "label_color": "1565C0"},
    "TIP": {"icon": "✦", "fill": "E8F5F2", "label_fill": "D4EFEB", "label_color": "00838F"},
    "WARNING": {"icon": "⚠", "fill": "FFF4E5", "label_fill": "FFE0B2", "label_color": "EF6C00"},
    "API": {"icon": "⬡", "fill": "F4ECFF", "label_fill": "E1D5F7", "label_color": "6A1B9A"},
    "DANGER": {"icon": "✖", "fill": "FFEBEE", "label_fill": "FFCDD2", "label_color": "C62828"},
}
A4_WIDTH_CM = 21.0
A4_HEIGHT_CM = 29.7
PORTRAIT_IMAGE_MAX_WIDTH_IN = 6.0
PORTRAIT_IMAGE_MAX_HEIGHT_IN = 8.2


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build the detailed MFDStudio user guide DOCX.")
    parser.add_argument("--source", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--rendered-diagram-dir", required=True)
    parser.add_argument("--skip-diagram-rendering", action="store_true")
    return parser.parse_args()


def resolve_path(repo_root: Path, raw_path: str) -> Path:
    path = Path(raw_path)
    if path.is_absolute():
        return path
    return (repo_root / path).resolve()


def resolve_document_image_path(source_dir: Path, relative_image: str, rendered_dir: Path) -> Path:
    normalized = relative_image.replace("\\", "/")
    if normalized.endswith(".png") and "/docs/user_guide/rendered/" in f"/{normalized}":
        return (rendered_dir / Path(normalized).name).resolve()
    return (source_dir / relative_image).resolve()


def render_diagrams(repo_root: Path, output_dir: Path) -> None:
    diagram_dir = repo_root / "docs" / "user_guide" / "diagrams"
    output_dir.mkdir(parents=True, exist_ok=True)
    for diagram_file in sorted(diagram_dir.glob("*.puml")):
        request = urllib.request.Request(
            url="https://kroki.io/plantuml/png",
            data=diagram_file.read_text(encoding="utf-8").encode("utf-8"),
            headers={"Content-Type": "text/plain"},
            method="POST",
        )
        output_file = output_dir / f"{diagram_file.stem}.png"
        with urllib.request.urlopen(request) as response:
            output_file.write_bytes(response.read())


def set_cell_shading(cell, fill: str) -> None:
    cell._tc.get_or_add_tcPr().append(parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{fill}"/>'))


def set_run_shading(run, fill: str) -> None:
    run._r.get_or_add_rPr().append(parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{fill}"/>'))


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def strip_inline_markdown(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def hex_to_rgb(hex_color: str) -> RGBColor:
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def chapter_group_hex(text: str) -> str:
    if text.startswith("Appendix "):
        return CHAPTER_GROUP_COLORS["teal"]

    match = re.match(r"^(\d+)", text)
    if not match:
        return CHAPTER_GROUP_COLORS["blue"]

    chapter = int(match.group(1))
    if 1 <= chapter <= 4:
        return CHAPTER_GROUP_COLORS["blue"]
    if 5 <= chapter <= 6:
        return CHAPTER_GROUP_COLORS["purple"]
    if chapter == 7:
        return CHAPTER_GROUP_COLORS["orange"]
    if 8 <= chapter <= 9:
        return CHAPTER_GROUP_COLORS["teal"]
    return CHAPTER_GROUP_COLORS["blue"]


def set_paragraph_bottom_border(paragraph, color: str, size: int = 12) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def append_inline_runs(
    paragraph,
    text: str,
    *,
    font_name: str = "Aptos",
    font_size: float = 10.0,
    font_color: RGBColor = BODY_TEXT_COLOR,
) -> None:
    for part in re.split(r"(`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            set_run_shading(run, "F5F5F5")
        else:
            clean = part.replace("**", "")
            run = paragraph.add_run(clean)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.color.rgb = font_color


def normalize_callout_label(label: str) -> str:
    lowered = label.strip().lower()
    if "danger" in lowered:
        return "DANGER"
    if "api" in lowered or "integration note" in lowered:
        return "API"
    if "tip" in lowered or "recommendation" in lowered:
        return "TIP"
    if "warning" in lowered or "important" in lowered or "common mistake" in lowered:
        return "WARNING"
    return "NOTE"


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    placeholder = paragraph.add_run("The table of contents will be updated by Word.")
    placeholder.italic = True
    run._r.append(fld_end)


def add_page_number(paragraph) -> None:
    prefix = paragraph.add_run("p. ")
    prefix.font.name = "Aptos"
    prefix.font.size = Pt(8.5)
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(doc: Document, title: str) -> None:
    section = doc.sections[0]
    configure_section_page(section, landscape=False)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.0)
    normal.font.color.rgb = BODY_TEXT_COLOR
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = Pt(15)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST

    title_style = doc.styles["Title"]
    title_style.font.name = "Aptos Display"
    title_style.font.size = Pt(24)
    title_style.font.color.rgb = RGBColor(0x2A, 0x4D, 0x69)

    for style_name, size, color in (
        ("Heading 1", 18, RGBColor(0x2A, 0x4D, 0x69)),
        ("Heading 2", 13.5, RGBColor(0x35, 0x50, 0x70)),
        ("Heading 3", 11.5, RGBColor(0x5C, 0x67, 0x7D)),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.page_break_before = False

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run(title)
    header_run.font.name = "Aptos"
    header_run.font.size = Pt(8.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)


def configure_section_page(section, *, landscape: bool) -> None:
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.4)
    if landscape:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = Cm(A4_HEIGHT_CM)
        section.page_height = Cm(A4_WIDTH_CM)
    else:
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = Cm(A4_WIDTH_CM)
        section.page_height = Cm(A4_HEIGHT_CM)


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    title_paragraph = doc.add_paragraph(style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_paragraph.add_run(title)
    title_run.font.size = Pt(28)

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.name = "Courier New"
    subtitle_run.font.size = Pt(12)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta_table = doc.add_table(rows=1, cols=1)
    meta_cell = meta_table.cell(0, 0)
    set_cell_shading(meta_cell, "EAF4F4")
    set_cell_margins(meta_cell)
    paragraph = meta_cell.paragraphs[0]
    paragraph.add_run("Document generated on ").bold = True
    paragraph.add_run(dt.datetime.now().strftime("%Y-%m-%d %H:%M"))
    paragraph.add_run("\nScope: ").bold = True
    paragraph.add_run("editor, architecture, generator, client API, offscreen runtime embedding, framebuffer plugin, launch scripts")
    paragraph.add_run("\nTarget audience: ").bold = True
    paragraph.add_run("MFD authors, C++ integrators, embedding application teams, and runtime teams")

    band_table = doc.add_table(rows=1, cols=3)
    for index, (label, fill) in enumerate((
        ("Editor · ch. 1-4", CHAPTER_GROUP_COLORS["blue"]),
        ("API · ch. 5-6", CHAPTER_GROUP_COLORS["purple"]),
        ("Plugin · ch. 7", CHAPTER_GROUP_COLORS["orange"]),
    )):
        cell = band_table.cell(0, index)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=100, start=100, bottom=100, end=100)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label)
        run.font.name = "Aptos"
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int) -> None:
    clean_text = strip_inline_markdown(text)
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    accent_hex = chapter_group_hex(clean_text)
    accent_rgb = hex_to_rgb(accent_hex)

    if level == 1:
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(10)
        run = paragraph.add_run(clean_text)
        run.font.name = "Aptos Display"
        run.font.size = Pt(18)
        run.font.color.rgb = accent_rgb
        set_paragraph_bottom_border(paragraph, accent_hex)
        return paragraph

    if level == 2:
        paragraph.paragraph_format.space_before = Pt(23)
        paragraph.paragraph_format.space_after = Pt(11)
        bullet = paragraph.add_run("▪ ")
        bullet.font.name = "Aptos Display"
        bullet.font.size = Pt(13.5)
        bullet.font.color.rgb = accent_rgb
        run = paragraph.add_run(clean_text)
        run.font.name = "Aptos Display"
        run.font.size = Pt(13.5)
        run.font.color.rgb = RGBColor(0x28, 0x35, 0x48)
        return paragraph

    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(clean_text)
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)
    return paragraph


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = Pt(15)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    append_inline_runs(paragraph, text)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing = Pt(15)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    append_inline_runs(paragraph, text)


def add_numbered(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.line_spacing = Pt(15)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    append_inline_runs(paragraph, text)


def add_code_block(doc: Document, lines: list[str], language: str) -> None:
    header = doc.add_paragraph()
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    header_run = header.add_run(f"CODE    {(language or 'text').upper()}")
    header_run.font.name = "Aptos"
    header_run.font.size = Pt(7.5)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor(0x35, 0x50, 0x70)
    set_run_shading(header_run, "DCE6F2")

    code_paragraph = doc.add_paragraph()
    code_paragraph.paragraph_format.space_before = Pt(0)
    code_paragraph.paragraph_format.space_after = Pt(8)
    code_paragraph.paragraph_format.left_indent = Cm(0.3)
    code_paragraph.paragraph_format.right_indent = Cm(0.3)
    code_paragraph.paragraph_format.line_spacing = Pt(12)
    code_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    code_run = code_paragraph.add_run("\n".join(lines))
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(8.6)
    code_run.font.color.rgb = BODY_TEXT_COLOR
    set_run_shading(code_run, "F1F5F9")


def add_note_block(doc: Document, lines: list[str]) -> None:
    label = "NOTE"
    body_lines = lines[:]
    if body_lines and body_lines[0].endswith(":"):
        label = normalize_callout_label(body_lines[0][:-1])
        body_lines = body_lines[1:]

    style = CALLOUT_STYLES[label]
    table = doc.add_table(rows=1, cols=2)
    label_cell = table.cell(0, 0)
    body_cell = table.cell(0, 1)
    set_cell_shading(label_cell, style["label_fill"])
    set_cell_shading(body_cell, style["fill"])
    set_cell_margins(label_cell, top=90, start=90, bottom=90, end=90)
    set_cell_margins(body_cell, top=90, start=140, bottom=90, end=120)

    label_paragraph = label_cell.paragraphs[0]
    label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = label_paragraph.add_run(f"{style['icon']} {label}")
    label_run.font.name = "Aptos"
    label_run.font.size = Pt(7.5)
    label_run.font.bold = True
    label_run.font.color.rgb = hex_to_rgb(style["label_color"])

    body_paragraph = body_cell.paragraphs[0]
    body_paragraph.paragraph_format.line_spacing = Pt(13)
    body_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    if not body_lines:
        body_lines = [label.title()]
    for index, line in enumerate(body_lines):
        if index > 0:
            body_paragraph.add_run().add_break()
        append_inline_runs(body_paragraph, line, font_size=9.0)


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    if len(lines) < 2:
        for line in lines:
            add_body_paragraph(doc, line)
        return

    def split_row(line: str) -> list[str]:
        trimmed = line.strip()
        if trimmed.startswith("|"):
            trimmed = trimmed[1:]
        if trimmed.endswith("|"):
            trimmed = trimmed[:-1]
        return [strip_inline_markdown(cell.strip()) for cell in trimmed.split("|")]

    header = split_row(lines[0])
    rows = [split_row(line) for line in lines[2:]]
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"

    for idx, value in enumerate(header):
        cell = table.cell(0, idx)
        cell.text = value
        set_cell_shading(cell, "1565C0" if idx == 0 else "1E88E5")
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Aptos"
            run.font.size = Pt(8.5)
            run.font.color.rgb = TABLE_HEADER_TEXT_COLOR

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            set_cell_shading(cell, "FFFFFF" if row_idx % 2 == 1 else "F8FAFF")
            set_cell_margins(cell)
            for run in cell.paragraphs[0].runs:
                run.font.name = "Aptos"
                run.font.size = Pt(9.5)
                run.font.color.rgb = BODY_TEXT_COLOR
                if col_idx == 0:
                    run.bold = True


def add_image(
    doc: Document,
    image_path: Path,
    caption: str,
    *,
    max_width_inches: float = PORTRAIT_IMAGE_MAX_WIDTH_IN,
    max_height_inches: float = PORTRAIT_IMAGE_MAX_HEIGHT_IN,
) -> None:
    with Image.open(image_path) as image:
        width_px, height_px = image.size

    scale = min(max_width_inches / width_px, max_height_inches / height_px)
    render_width_inches = width_px * scale
    render_height_inches = height_px * scale

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(
        str(image_path),
        width=Inches(render_width_inches),
        height=Inches(render_height_inches),
    )
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(0)
    caption_paragraph.paragraph_format.space_after = Pt(8)
    caption_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    caption_paragraph.paragraph_format.keep_together = True
    caption_run = caption_paragraph.add_run(caption)
    caption_run.font.name = "Aptos"
    caption_run.font.size = Pt(9)
    caption_run.italic = True


def build_document(source_path: Path, output_path: Path, repo_root: Path, rendered_dir: Path) -> None:
    lines = source_path.read_text(encoding="utf-8").splitlines()
    title = next((line[2:].strip() for line in lines if line.startswith("# ")), "MFDStudio - Detailed User Guide")
    subtitle = next((line.split(":", 1)[1].strip() for line in lines if line.startswith("Subtitle:")), "")

    doc = Document()
    configure_document(doc, title)
    add_cover(doc, title, subtitle)

    paragraph_buffer: list[str] = []
    table_buffer: list[str] = []
    code_buffer: list[str] = []
    note_buffer: list[str] = []
    in_code_block = False
    code_fence = "```"
    code_language = ""
    source_dir = source_path.parent
    pending_page_break = False

    def flush_paragraph() -> None:
        if paragraph_buffer:
            add_body_paragraph(doc, " ".join(part.strip() for part in paragraph_buffer if part.strip()))
            paragraph_buffer.clear()

    def flush_table() -> None:
        if table_buffer:
            add_markdown_table(doc, table_buffer[:])
            table_buffer.clear()

    def flush_code() -> None:
        if code_buffer:
            add_code_block(doc, code_buffer[:], code_language)
            code_buffer.clear()

    def flush_note() -> None:
        if note_buffer:
            add_note_block(doc, note_buffer[:])
            note_buffer.clear()

    def apply_pending_page_break(paragraph) -> None:
        nonlocal pending_page_break
        if pending_page_break:
            paragraph.paragraph_format.page_break_before = True
            pending_page_break = False

    for index, line in enumerate(lines):
        if index == 0 and line.startswith("# "):
            continue
        if line.startswith("Subtitle:"):
            continue

        trimmed = line.strip()

        if in_code_block:
            if trimmed.startswith(code_fence):
                flush_code()
                in_code_block = False
            else:
                code_buffer.append(line)
            continue

        if trimmed.startswith(code_fence):
            flush_paragraph()
            flush_table()
            flush_note()
            code_language = trimmed[len(code_fence):].strip()
            in_code_block = True
            continue

        if trimmed.startswith("![") and "](" in trimmed and trimmed.endswith(")"):
            flush_paragraph()
            flush_table()
            flush_note()
            separator_index = trimmed.index("](")
            caption = trimmed[2:separator_index].strip()
            relative_image = trimmed[separator_index + 2:-1].strip()
            image_path = resolve_document_image_path(source_dir, relative_image, rendered_dir)

            add_image(doc, image_path, caption)
            continue

        if trimmed.startswith("|"):
            flush_paragraph()
            flush_note()
            table_buffer.append(line)
            continue

        if table_buffer and not trimmed.startswith("|"):
            flush_table()

        if trimmed.startswith("> "):
            flush_paragraph()
            note_buffer.append(trimmed[2:].strip())
            continue

        if note_buffer and not trimmed.startswith("> "):
            flush_note()

        if not trimmed:
            flush_paragraph()
            flush_note()
            continue

        if trimmed == "[[TOC]]":
            _ = doc.add_heading("Table of Contents", level=1)
            add_toc(doc.add_paragraph())
            continue

        if trimmed == "<!-- PAGEBREAK -->":
            flush_paragraph()
            flush_note()
            pending_page_break = True
            continue

        if trimmed.startswith("### "):
            flush_paragraph()
            flush_note()
            apply_pending_page_break(add_heading(doc, trimmed[4:].strip(), 3))
            continue

        if trimmed.startswith("## "):
            flush_paragraph()
            flush_note()
            apply_pending_page_break(add_heading(doc, trimmed[3:].strip(), 2))
            continue

        if trimmed.startswith("# "):
            flush_paragraph()
            flush_note()
            apply_pending_page_break(add_heading(doc, trimmed[2:].strip(), 1))
            continue

        if re.match(r"^\d+\.\s", trimmed):
            flush_paragraph()
            flush_note()
            add_numbered(doc, trimmed)
            continue

        if trimmed.startswith("- "):
            flush_paragraph()
            flush_note()
            add_bullet(doc, trimmed[2:].strip())
            continue

        paragraph_buffer.append(trimmed)

    flush_paragraph()
    flush_table()
    flush_note()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> int:
    args = parse_args()
    repo_root = Path(__file__).resolve().parents[2]
    source_path = resolve_path(repo_root, args.source)
    output_path = resolve_path(repo_root, args.output)
    rendered_dir = resolve_path(repo_root, args.rendered_diagram_dir)

    if not args.skip_diagram_rendering:
        render_diagrams(repo_root, rendered_dir)

    build_document(source_path, output_path, repo_root, rendered_dir)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
